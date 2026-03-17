"""
File export utilities – generate PDF and DOCX brand kits.
"""
import os
import json
from io import BytesIO
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
    PageBreak,
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)


# ── PDF Generation ────────────────────────────────────────────────────

def _add_section(elements: list, title: str, body: str, styles) -> None:
    elements.append(Paragraph(title, styles["Heading2"]))
    elements.append(Spacer(1, 6))
    for line in body.strip().split("\n"):
        elements.append(Paragraph(line, styles["BodyText"]))
        elements.append(Spacer(1, 3))
    elements.append(Spacer(1, 12))


def generate_pdf(project_id: str, brand_data: dict) -> str:
    """Create a brand kit PDF and return the file path."""
    filepath = EXPORT_DIR / f"{project_id}_brand_kit.pdf"

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="BrandTitle",
        fontSize=24,
        leading=30,
        textColor=HexColor("#1a1a2e"),
        spaceAfter=20,
        alignment=1,  # center
    ))

    elements: list = []

    # Title page
    elements.append(Spacer(1, 2 * inch))
    brand_name = brand_data.get("naming", {}).get("brand_name", "Brand Kit")
    elements.append(Paragraph(f"🎨 {brand_name} – Brand Identity Kit", styles["BrandTitle"]))
    elements.append(Spacer(1, 0.5 * inch))
    tagline = brand_data.get("naming", {}).get("tagline", "")
    if tagline:
        elements.append(Paragraph(tagline, styles["BodyText"]))
    elements.append(PageBreak())

    # Sections
    section_map = {
        "idea_discovery": "💡 Idea Discovery",
        "market_research": "📊 Market Research",
        "competitor_analysis": "🔍 Competitor Analysis",
        "brand_strategy": "🎯 Brand Strategy",
        "naming": "✏️ Brand Naming",
        "design_agent": "🎨 Design Direction",
        "content_agent": "📝 Brand Content",
        "guidelines_agent": "📋 Brand Guidelines",
    }

    for key, title in section_map.items():
        section = brand_data.get(key)
        if section:
            body = json.dumps(section, indent=2, default=str)
            _add_section(elements, title, body, styles)

    doc.build(elements)
    return str(filepath)


# ── DOCX Generation ──────────────────────────────────────────────────

def generate_docx(project_id: str, brand_data: dict) -> str:
    """Create a brand kit DOCX and return the file path."""
    filepath = EXPORT_DIR / f"{project_id}_brand_kit.docx"

    doc = Document()

    # Title
    brand_name = brand_data.get("naming", {}).get("brand_name", "Brand Kit")
    title = doc.add_heading(f"{brand_name} – Brand Identity Kit", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline = brand_data.get("naming", {}).get("tagline", "")
    if tagline:
        p = doc.add_paragraph(tagline)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    section_map = {
        "idea_discovery": "Idea Discovery",
        "market_research": "Market Research",
        "competitor_analysis": "Competitor Analysis",
        "brand_strategy": "Brand Strategy",
        "naming": "Brand Naming",
        "design_agent": "Design Direction",
        "content_agent": "Brand Content",
        "guidelines_agent": "Brand Guidelines",
    }

    for key, title_text in section_map.items():
        section = brand_data.get(key)
        if section:
            doc.add_heading(title_text, level=1)
            content = json.dumps(section, indent=2, default=str)
            for line in content.split("\n"):
                doc.add_paragraph(line, style="List Bullet")

    doc.save(str(filepath))
    return str(filepath)
